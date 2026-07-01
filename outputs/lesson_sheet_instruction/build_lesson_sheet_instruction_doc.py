from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("outputs/lesson_sheet_instruction")
OUT_PATH = OUT_DIR / "lesson_sheet_data_entry_instruction.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B8C3D1"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None):
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rfonts.set(qn("w:cs"), name)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))


def set_col_widths(table, widths_inches):
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.first_child_found_in("w:tcW")
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:type"), "dxa")
            tcw.set(qn("w:w"), str(int(width * 1440)))


def format_table(table, header_fill=LIGHT_BLUE, widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    set_table_width(table)
    if widths:
        set_col_widths(table, widths)

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(run, size=9.5 if row_idx == 0 else 9)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, size=9.5, bold=True, color=DARK_BLUE)


def add_cell_text(cell, text, bold=False, color=None):
    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    run = paragraph.add_run(text)
    set_run_font(run, size=9, bold=bold, color=color)


def add_paragraph(doc, text="", style=None, before=None, after=None, line=1.25, color=None, bold=None):
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, color=color, bold=bold)
    paragraph.paragraph_format.space_before = Pt(0 if before is None else before)
    paragraph.paragraph_format.space_after = Pt(6 if after is None else after)
    paragraph.paragraph_format.line_spacing = line
    return paragraph


def add_code_paragraph(doc, label, value):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    run = paragraph.add_run(f"{label}: ")
    set_run_font(run, size=10, bold=True, color=DARK_BLUE)
    code = paragraph.add_run(value)
    set_run_font(code, name="Consolas", size=9.5)
    return paragraph


def add_note(doc, title, body, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    format_table(table, header_fill=fill, widths=[6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.text = ""
    run = p.add_run(title)
    set_run_font(run, size=10, bold=True, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(3)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    run2 = p2.add_run(body)
    set_run_font(run2, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(level=level)
    paragraph.text = ""
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
        paragraph.paragraph_format.space_before = Pt(18)
        paragraph.paragraph_format.space_after = Pt(10)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
        paragraph.paragraph_format.space_before = Pt(14)
        paragraph.paragraph_format.space_after = Pt(7)
    else:
        set_run_font(run, size=12, color=DARK_BLUE, bold=True)
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_two_col_table(doc, rows, widths=(1.9, 4.6), header=None):
    row_count = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=row_count, cols=2)
    start = 0
    if header:
        add_cell_text(table.cell(0, 0), header[0], bold=True)
        add_cell_text(table.cell(0, 1), header[1], bold=True)
        start = 1
    for idx, (left, right) in enumerate(rows, start=start):
        add_cell_text(table.cell(idx, 0), left, bold=True, color=DARK_BLUE)
        add_cell_text(table.cell(idx, 1), right)
    format_table(table, widths=list(widths))
    return table


def add_three_col_table(doc, rows, widths=(1.55, 2.25, 2.7), header=("Колонка", "Когда заполнять", "Что писать")):
    table = doc.add_table(rows=len(rows) + 1, cols=3)
    for i, h in enumerate(header):
        add_cell_text(table.cell(0, i), h, bold=True)
    for r, row in enumerate(rows, start=1):
        for c, text in enumerate(row):
            add_cell_text(table.cell(r, c), text, bold=(c == 0), color=DARK_BLUE if c == 0 else None)
    format_table(table, widths=list(widths))
    return table


def set_document_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        for element in style.element.xpath(".//w:rFonts"):
            element.set(qn("w:ascii"), "Calibri")
            element.set(qn("w:hAnsi"), "Calibri")
            element.set(qn("w:eastAsia"), "Microsoft YaHei")
            element.set(qn("w:cs"), "Calibri")


def add_footer(doc):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.text = ""
    run = paragraph.add_run("RussianBridge - инструкция по заполнению таблицы уроков")
    set_run_font(run, size=8.5, color=MUTED)


def build_document():
    doc = Document()
    set_document_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.paragraph_format.line_spacing = 1.0
    run = title.add_run("Инструкция по заполнению Google-таблицы уроков RussianBridge")
    set_run_font(run, size=22, color=DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Памятка для редактора данных: как добавлять уроки, упражнения, карточки, варианты ответов и медиа без работы с кодом.")
    set_run_font(run, size=11, color=MUTED)

    add_note(
        doc,
        "Главная идея",
        "В таблице одна строка означает один вопрос или одну карточку. Строки с одинаковыми lesson_number и exercise_order собираются приложением в одно упражнение. Поэтому порядок и обязательные поля важнее внешнего оформления таблицы.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "1. Перед началом", 1)
    add_bullet(doc, "Работайте в Google Sheets, а не в выгруженном Excel-файле, если задача - обновить данные для приложения.")
    add_bullet(doc, "Не удаляйте и не переименовывайте первую строку с названиями колонок.")
    add_bullet(doc, "Не объединяйте ячейки и не вставляйте картинки прямо в ячейку. Для картинок и аудио пишется ссылка или путь к файлу.")
    add_bullet(doc, "Если нужно добавить похожую строку, сначала скопируйте существующую строку того же типа упражнения, затем замените значения.")
    add_bullet(doc, "Форматирование ячейки в Google Sheets не попадает в приложение. Если в тексте нужно выделение жирным, используйте теги <b> и </b>, как в существующих примерах.")

    add_heading(doc, "2. Как устроена таблица", 1)
    add_paragraph(
        doc,
        "Предпочтительный вариант - отдельный лист на каждый урок. Листы называются Lesson1, Lesson2, Lesson3 и так далее. Служебные листы вроде Summary и README можно оставить: приложение их пропускает.",
    )
    add_two_col_table(
        doc,
        [
            ("Lesson1, Lesson2...", "Рабочие листы уроков. В них вносятся строки с упражнениями."),
            ("Summary, README", "Служебные листы. Их можно читать, но не нужно заполнять как уроки."),
            ("Первая строка", "Заголовки колонок. Ее нельзя менять."),
            ("Одна строка", "Одна карточка, один вопрос или один вариант задания внутри упражнения."),
        ],
        header=("Элемент таблицы", "Что это значит"),
    )

    add_heading(doc, "3. Обязательные колонки", 1)
    add_paragraph(doc, "Эти поля заполняются в каждой рабочей строке. Если одно из них пустое, строка может быть пропущена приложением.")
    add_three_col_table(
        doc,
        [
            ("lesson_number", "Всегда", "Номер урока: 1, 2, 3... Должен совпадать с листом Lesson1, Lesson2 и т.д."),
            ("lesson_name", "Всегда", "Название урока, которое видно пользователю, например Урок 5."),
            ("exercise_order", "Всегда", "Порядковый номер упражнения внутри урока: 1, 2, 3..."),
            ("exercise_type", "Всегда", "Один из вариантов: Flashcards, Translate, FillBlank, MakeSentence, Writing."),
        ],
    )
    add_note(
        doc,
        "Важно про порядок",
        "Если несколько строк имеют одинаковые lesson_number и exercise_order, они относятся к одному упражнению. У всех таких строк должен быть один и тот же exercise_type.",
    )

    add_heading(doc, "4. Типы упражнений", 1)
    add_paragraph(doc, "Ниже указаны поля, которые нужно заполнять для каждого типа. Остальные колонки в этой строке можно оставить пустыми.")

    add_heading(doc, "Flashcards", 2)
    add_paragraph(doc, "Карточки для изучения слов. Обычно несколько строк с одним exercise_order образуют одну колоду карточек.")
    add_three_col_table(
        doc,
        [
            ("foreign_word", "Обязательно", "Слово или фраза на китайском/иностранном языке."),
            ("translation", "Обязательно", "Перевод на русский."),
            ("example_foreign", "Желательно", "Пример предложения на иностранном языке."),
            ("example_translation", "Желательно", "Русский перевод примера."),
            ("image", "По необходимости", "Путь или ссылка на картинку."),
            ("front_audio", "По необходимости", "Аудио для лицевой стороны карточки."),
            ("back_audio", "По необходимости", "Аудио для обратной стороны карточки."),
            ("is_grammar_cards", "Редко", "Пишите true только для грамматических карточек. В остальных случаях false или пусто."),
        ],
    )

    add_heading(doc, "Translate", 2)
    add_paragraph(doc, "Вопрос с выбором правильного перевода или варианта.")
    add_three_col_table(
        doc,
        [
            ("foreign_word", "Обязательно", "Текст вопроса или слово, которое нужно перевести."),
            ("correct_translation", "Обязательно", "Правильный ответ."),
            ("options", "Обязательно", "Все варианты ответа через вертикальную черту: вариант1|вариант2|вариант3."),
            ("image", "По необходимости", "Картинка к вопросу."),
            ("audio", "По необходимости", "Аудио к вопросу."),
        ],
    )

    add_heading(doc, "FillBlank", 2)
    add_paragraph(doc, "Упражнение, где ученик вставляет пропущенное слово.")
    add_three_col_table(
        doc,
        [
            ("task_title", "Желательно", "Короткая инструкция, например Выберите правильную форму."),
            ("hint", "По необходимости", "Подсказка, правило или грамматический комментарий."),
            ("sentence_with_blanks", "Обязательно", "Предложение с пропуском. Для пропуска используйте ___."),
            ("correct_answers", "Обязательно", "Правильный ответ. Если вариантов несколько, разделите их через |."),
            ("word_bank", "Желательно", "Слова для выбора через |."),
        ],
    )

    add_heading(doc, "MakeSentence", 2)
    add_paragraph(doc, "Упражнение, где ученик собирает предложение из слов.")
    add_three_col_table(
        doc,
        [
            ("task_title", "Желательно", "Инструкция, например Составьте предложение."),
            ("hint", "По необходимости", "Подсказка."),
            ("shuffled_words", "Обязательно", "Слова в перемешанном порядке через |."),
            ("correct_sentences", "Обязательно", "Правильное предложение. Если допустимы несколько вариантов, разделите их через |."),
        ],
    )

    add_heading(doc, "Writing", 2)
    add_paragraph(doc, "Упражнение на написание слов, обычно с аудио.")
    add_three_col_table(
        doc,
        [
            ("correct_words", "Обязательно", "Слова, которые должен написать ученик, через |."),
            ("audio", "Желательно", "Аудио для слов. Если аудио несколько, разделите пути через | в том же порядке, что и correct_words."),
        ],
    )

    add_heading(doc, "5. Как добавлять данные", 1)
    add_heading(doc, "Добавить новую карточку в Flashcards", 2)
    add_number(doc, "Откройте лист нужного урока, например Lesson5.")
    add_number(doc, "Найдите существующую строку с exercise_type = Flashcards.")
    add_number(doc, "Скопируйте строку ниже последней карточки той же колоды.")
    add_number(doc, "Оставьте lesson_number, lesson_name, exercise_order и exercise_type такими же.")
    add_number(doc, "Замените foreign_word, translation и примеры на новые данные.")

    add_heading(doc, "Добавить новое упражнение в урок", 2)
    add_number(doc, "Найдите последний exercise_order на листе урока.")
    add_number(doc, "Создайте новую строку и поставьте следующий номер exercise_order.")
    add_number(doc, "Выберите exercise_type: Flashcards, Translate, FillBlank, MakeSentence или Writing.")
    add_number(doc, "Заполните обязательные поля для выбранного типа упражнения.")
    add_number(doc, "Если упражнение состоит из нескольких вопросов, добавьте несколько строк с тем же exercise_order и тем же exercise_type.")

    add_heading(doc, "Добавить новый урок", 2)
    add_number(doc, "Скопируйте лист предыдущего урока.")
    add_number(doc, "Переименуйте копию в LessonN, где N - новый номер урока, например Lesson11.")
    add_number(doc, "Во всех рабочих строках замените lesson_number на новый номер.")
    add_number(doc, "Замените lesson_name на название нового урока.")
    add_number(doc, "Удалите старые строки урока и внесите новые упражнения, сохранив первую строку с заголовками.")

    add_heading(doc, "6. Списки, варианты ответов и медиа", 1)
    add_paragraph(doc, "В одной ячейке иногда нужно перечислить несколько значений. Для этого используется вертикальная черта |.")
    add_code_paragraph(doc, "options", "идти|ехать|бежать|стоять")
    add_code_paragraph(doc, "word_bank", "архитектор|архитектором|архитектора|архитекторе")
    add_code_paragraph(doc, "correct_sentences", "Мы пойдём в галерею в субботу|В субботу мы пойдём в галерею")
    add_bullet(doc, "Не разделяйте варианты запятыми, если это именно список для приложения.")
    add_bullet(doc, "Не ставьте вертикальную черту внутри обычного текста, если она не должна разделять варианты.")
    add_bullet(doc, "Пробелы вокруг | лучше не ставить: приложение их обычно обрежет, но без пробелов проще проверять.")

    add_heading(doc, "Медиа", 2)
    add_paragraph(doc, "В ячейки image, audio, front_audio и back_audio пишется путь или ссылка. Сам файл в таблицу не вставляется.")
    add_two_col_table(
        doc,
        [
            ("URL", "https://example.com/card1.png"),
            ("Resources-путь", "Sprites/Lesson/Lesson5/L5_Ex4_WalksInPark"),
            ("Путь проекта", "Assets/Sprites/Lesson/Lesson5/L5_Ex4_WalksInPark.png"),
        ],
        widths=(1.55, 4.95),
        header=("Тип", "Пример"),
    )
    add_note(
        doc,
        "Если сомневаетесь в медиа",
        "Для новых картинок и аудио лучше попросить разработчика дать готовый путь. В приложении особенно важно, чтобы файл был доступен в сборке, а не только лежал на компьютере редактора.",
    )

    add_heading(doc, "7. Служебные колонки", 1)
    add_paragraph(doc, "В выгрузке могут быть колонки source_asset, source_question_index и source_media_notes. Они нужны для внутренней проверки и истории импорта.")
    add_bullet(doc, "Для обычного ручного ввода новых строк эти колонки можно оставить пустыми.")
    add_bullet(doc, "Не используйте source_media_notes как место для комментариев к уроку: приложение эту информацию не показывает ученику.")
    add_bullet(doc, "Если в source_media_notes написано, что путь к медиа вне Assets/Resources, это сигнал показать строку разработчику.")

    add_heading(doc, "8. Примеры заполнения", 1)
    add_heading(doc, "Пример Flashcards", 2)
    add_code_paragraph(doc, "lesson_number", "1")
    add_code_paragraph(doc, "exercise_order", "1")
    add_code_paragraph(doc, "exercise_type", "Flashcards")
    add_code_paragraph(doc, "foreign_word", "申请表")
    add_code_paragraph(doc, "translation", "анкета")
    add_code_paragraph(doc, "example_foreign", "在学校，老师给了我们一张<b>申请表</b>。")
    add_code_paragraph(doc, "example_translation", "В школе учитель дал нам <b>анкету</b>.")

    add_heading(doc, "Пример Translate", 2)
    add_code_paragraph(doc, "exercise_type", "Translate")
    add_code_paragraph(doc, "foreign_word", "请选择正确的完成体动词 «ехать».")
    add_code_paragraph(doc, "correct_translation", "поехать")
    add_code_paragraph(doc, "options", "поехать|ездить|есть|поесть")

    add_heading(doc, "Пример FillBlank", 2)
    add_code_paragraph(doc, "exercise_type", "FillBlank")
    add_code_paragraph(doc, "task_title", "请选择")
    add_code_paragraph(doc, "hint", "работать + пятый падеж")
    add_code_paragraph(doc, "sentence_with_blanks", "Его профессия - преподаватель, но он работает ___")
    add_code_paragraph(doc, "correct_answers", "архитектором")
    add_code_paragraph(doc, "word_bank", "архитектор|архитектором|архитектора|архитекторе")

    add_heading(doc, "Пример MakeSentence", 2)
    add_code_paragraph(doc, "exercise_type", "MakeSentence")
    add_code_paragraph(doc, "task_title", "请用下面的词造句:")
    add_code_paragraph(doc, "shuffled_words", "галерею|мы|в|пойдём|субботу")
    add_code_paragraph(doc, "correct_sentences", "Мы пойдём в галерею в субботу|В субботу мы пойдём в галерею")

    add_heading(doc, "Пример Writing", 2)
    add_code_paragraph(doc, "exercise_type", "Writing")
    add_code_paragraph(doc, "correct_words", "привет|меня|зовут")
    add_code_paragraph(doc, "audio", "Audio/Lesson1/privet|Audio/Lesson1/menya|Audio/Lesson1/zovut")

    add_heading(doc, "9. Чек-лист перед сдачей", 1)
    checklist = [
        "Первая строка с заголовками осталась на месте.",
        "На листе урока нет пустых lesson_number, exercise_order и exercise_type.",
        "exercise_type написан одним из пяти допустимых значений.",
        "Все строки одного упражнения имеют одинаковый exercise_order и exercise_type.",
        "Списки вариантов разделены символом |, а не запятыми.",
        "В FillBlank есть sentence_with_blanks и correct_answers.",
        "В Translate есть correct_translation и options, где правильный ответ тоже есть среди options.",
        "В Flashcards есть foreign_word и translation.",
        "Пути к картинкам и аудио не являются локальными файлами с компьютера редактора.",
        "Случайные пробелы в конце важных слов удалены, особенно в переводах и правильных ответах.",
    ]
    for item in checklist:
        add_bullet(doc, item)

    add_heading(doc, "10. Частые ошибки", 1)
    add_two_col_table(
        doc,
        [
            ("Изменили название колонки", "Верните старое название из первой строки или скопируйте его из соседнего листа."),
            ("Одинаковый exercise_order у разных типов", "Разделите их на разные номера упражнений."),
            ("Варианты написаны через запятую", "Замените запятые между вариантами на |."),
            ("Картинка вставлена в ячейку", "Удалите вставленную картинку и напишите путь или URL в колонке image."),
            ("Жирный текст сделан форматированием Google Sheets", "Используйте <b>текст</b>, если приложение должно увидеть выделение."),
            ("Новый лист назван по-русски", "Переименуйте лист в формате LessonN, например Lesson11."),
        ],
        widths=(2.15, 4.35),
        header=("Ошибка", "Как исправить"),
    )

    add_heading(doc, "Короткая формула работы", 1)
    add_note(
        doc,
        "Скопировать - заменить - проверить",
        "Самый безопасный способ заполнения: скопировать похожую готовую строку, заменить только содержательные поля, проверить обязательные колонки и списки через |. Это снижает риск сломать структуру таблицы.",
        fill=LIGHT_BLUE,
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = build_document()
    print(path.resolve())
